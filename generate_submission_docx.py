from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

# Initialize document
doc = Document()
doc.add_heading("From College Boards to Dashboards — Submission Document", 0)
doc.add_paragraph("Author: Pascal Esegemou Ekenya Fonjock")
doc.add_paragraph("Date: 2025-09-15")

# Executive Summary
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This story explores how professor effectiveness (EvaluationScore) relates to course costs and enrollment. "
    "We identify best practice professors who earn high evaluation scores at moderate or low cost, "
    "and low-cost courses with poor evaluations that would benefit most from targeted interventions."
)

# Helper function to add image from URL
def add_image_from_url(url, caption):
    response = requests.get(url)
    image_stream = BytesIO(response.content)
    doc.add_picture(image_stream, width=Inches(6))
    doc.add_paragraph(caption, style='Caption')

# Dataset Preparation Evidence
doc.add_heading("2. Dataset Preparation Evidence", level=1)
add_image_from_url("https://i.postimg.cc/3WzYjYPb/01-attempt-s3-manifest-error-png.png",
                   "Attempt to create dataset from fictional S3 manifest (expected error). Rubric: Section 1 evidence.")
add_image_from_url("https://i.postimg.cc/kVjd2mQs/03-dataset-list-q-student-enrollment-png.png",
                   "Q - Student Enrollment dataset visible in Datasets. Rubric: Section 1 evidence.")
add_image_from_url("https://i.postimg.cc/3yrTrRvW/04-dataset-refresh-schedule-png.png",
                   "Weekly refresh schedule set to Sunday 12:00 AM (timezone visible). Rubric: Section 1.")
add_image_from_url("https://i.postimg.cc/2bMRx7XB/06-rename-homeoforigin-to-nationalorigin-png.png",
                   "Field HomeOfOrigin renamed to NationalOrigin with description. Rubric: Section 1.")
add_image_from_url("https://i.postimg.cc/47mRtSwg/07-student-type-calculated-field-png.png",
                   "Calculated field Student Type defined. Rubric: Section 1.")

# Visuals Created Using Q
doc.add_heading("3. Visuals Created Using Q", level=1)
add_image_from_url("https://i.postimg.cc/YhcBG8W5/09-visual-student-majors-by-year-initial-png.png",
                   "Visual created by Q — Student Majors by Year (initial)")
add_image_from_url("https://i.postimg.cc/qN4fQRQQ/11-student-majors-by-year-vertical-bar-png.png",
                   "Visual reconfigured to vertical bar chart (formatted, no comma separators)")
add_image_from_url("https://i.postimg.cc/rdwBXPBr/12-proportion-of-student-types-pie-png.png",
                   "Proportion of Student Types (pie chart)")

# Topic & Named Entities
doc.add_heading("4. Topic & Named Entities", level=1)
add_image_from_url("https://i.postimg.cc/DmXtTsMB/16-topic-fields-list-png.png",
                   "Topic field list includes required fields.")
add_image_from_url("https://i.postimg.cc/21cHv8t1/17-topic-entities-student-details-png.png",
                   "Named Entity: Student Details.")
add_image_from_url("https://i.postimg.cc/hJX20RFp/18-topic-entities-course-details-png.png",
                   "Named Entity: Course Details.")
add_image_from_url("https://i.postimg.cc/qtWjbD40/19-topic-entities-prof-eval-png.png",
                   "Named Entity: Professor Evaluation.")
add_image_from_url("https://i.postimg.cc/HcG94J9k/20-q-best-instructors-verified-png.png",
                   "Verified Q answer: Best instructors.")

# Dashboard, Scenarios & Thread
doc.add_heading("5. Dashboard, Scenarios & Thread", level=1)
add_image_from_url("https://i.postimg.cc/xNnGhW7r/23-publish-dashboard-q-enabled-png.png",
                   "Student Enrollment Dashboard published with Q enabled.")
add_image_from_url("https://i.postimg.cc/3yGcTQLW/25-scenario-create-select-visuals-png.png",
                   "Scenario created with selected visuals.")
add_image_from_url("https://i.postimg.cc/yDtbt8yw/26-scenario-starter-png.png",
                   "Starter question: How do we improve professor evaluations while avoiding increased cost per course?")
add_image_from_url("https://i.postimg.cc/yDMLnrVK/34-scenario-thread-full-png.png",
                   "Scenario thread showing iterative Q.")

# Data Story (example visual)
doc.add_heading("6. Data Story", level=1)
add_image_from_url("https://i.postimg.cc/qN4fQRQQ/11-student-majors-by-year-vertical-bar-png.png",
                   "Data Story Visual: Student Majors by Year")
add_image_from_url("https://i.postimg.cc/rdwBXPBr/12-proportion-of-student-types-pie-png.png",
                   "Data Story Visual: Proportion of Student Types")
add_image_from_url("https://i.postimg.cc/dZsHrkJT/29-vis-courses-best-eval-png.png",
                   "Data Story Visual: Avg EvaluationScore by Course")
add_image_from_url("https://i.postimg.cc/t7kmy5zd/31-vis-courses-highest-cost-png.png",
                   "Data Story Visual: Avg CostPerCourse by Course")

# Other Sections
doc.add_heading("7. File listings and additional artifacts", level=1)
doc.add_paragraph("dataset_fields.md, calculated_fields.md, Q_prompts.md, verified_answers.md, "
                  "scenario_thread.md, data_story.md, manifest_sample.json, screenshots/ (all screenshot files present)")

doc.add_heading("8. Methodology", level=1)
doc.add_paragraph("Steps: sample dataset usage, calculated fields creation, Q prompts, topic creation, scenario & thread, data story creation.")

doc.add_heading("9. Originality Statement", level=1)
doc.add_paragraph("I confirm this submission is my original work. Any referenced AWS/Udacity documentation is cited.")

doc.add_heading("10. Checklist for Rubric", level=1)
doc.add_paragraph("Map each rubric criterion to the screenshot and file you provided.")

# Save DOCX
doc.save("submission_document.docx")
print("submission_document.docx has been created successfully!")
