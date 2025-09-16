cat > generate_submission_doc.py << 'PY'
#!/usr/bin/env python3
"""
Generate submission_document.docx using screenshots in ./screenshots.
This script is robust: if an expected image file is missing it inserts a note.
"""

import os
from docx import Document
from docx.shared import Inches

# Helper to find an image file with several common filename variants
def find_image_file(base_name):
    # base_name expected to include extension e.g. "01_attempt_s3_manifest_error.png"
    base_no_ext = base_name[:-4] if base_name.lower().endswith('.png') else base_name
    candidates = []

    # Candidate variations
    variants = [
        f"{base_no_ext}.png",
        f"{base_no_ext}.PNG",
        f"{base_no_ext.replace('_','-')}.png",
        f"{base_no_ext.replace('-','_')}.png",
        f"{base_no_ext.replace('_','-')}-png.png",
        f"{base_no_ext}-png.png"
    ]
    for v in variants:
        candidates.append(os.path.join("screenshots", v))

    # Also try to find file containing the numeric prefix (e.g., "01")
    prefix = base_no_ext.split('_')[0]
    try:
        for f in os.listdir("screenshots"):
            if f.lower().startswith(prefix.lower()) and f.lower().endswith('.png'):
                candidates.append(os.path.join("screenshots", f))
    except FileNotFoundError:
        # screenshots folder not present
        return None

    for c in candidates:
        if os.path.isfile(c):
            return c
    return None

def insert_image_or_note(doc, filename, caption, width_in_inches=6.5):
    path = find_image_file(filename)
    if path:
        try:
            doc.add_picture(path, width=Inches(width_in_inches))
            doc.add_paragraph(caption)
            print(f"Inserted image: {path}")
            return True
        except Exception as e:
            doc.add_paragraph(f"[ERROR inserting image {path}]: {e}")
            print(f"ERROR inserting {path}: {e}")
            return False
    else:
        doc.add_paragraph(f"[MISSING IMAGE] {filename} — {caption}")
        print(f"MISSING: {filename}")
        return False

def main():
    doc = Document()
    # Title
    doc.add_heading("From College Boards to Dashboards — Submission Document", 0)
    doc.add_paragraph("Author: Your Full Name")
    doc.add_paragraph("Date: 2025-09-15")
    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("In MS Word: References → Table of Contents → Update field to generate TOC.")
    doc.add_page_break()

    # 1. Executive Summary (try reading from data_story.md)
    doc.add_heading("1. Executive Summary", level=1)
    if os.path.isfile("data_story.md"):
        with open("data_story.md", "r", encoding="utf-8") as f:
            content = f.read().strip()
            if content:
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("Executive summary: (paste from data_story.md here).")
    else:
        doc.add_paragraph("Executive summary: (paste from data_story.md here).")

    # 2. Dataset Preparation Evidence
    doc.add_heading("2. Dataset Preparation Evidence", level=1)
    dataset_images = [
        ("01_attempt_s3_manifest_error.png", "Attempt to create dataset from fictional S3 manifest (expected error). Rubric: Section 1 evidence."),
        ("03_dataset_list_q_student_enrollment.png", "Q - Student Enrollment dataset visible in Datasets. Rubric: Section 1 evidence."),
        ("04_dataset_refresh_schedule.png", "Weekly refresh schedule set to Sunday 12:00 AM (timezone visible). Rubric: Section 1."),
        ("06_rename_homeoforigin_to_nationalorigin.png", "Field HomeOfOrigin renamed to NationalOrigin with description. Rubric: Section 1."),
        ("07_student_type_calculated_field.png", "Calculated field Student Type defined. Rubric: Section 1."),
    ]
    for fn, caption in dataset_images:
        insert_image_or_note(doc, fn, caption)

    # 3. Visuals Created Using Q
    doc.add_heading("3. Visuals Created Using Q", level=1)
    visuals = [
        ("09_visual_student_majors_by_year_initial.png", "Visual created by Q — Student Majors by Year (initial)"),
        ("11_student_majors_by_year_vertical_bar.png", "Visual reconfigured to vertical bar chart (formatted, no comma separators)"),
        ("12_proportion_of_student_types_pie.png", "Proportion of Student Types (pie chart)"),
    ]
    for fn, caption in visuals:
        insert_image_or_note(doc, fn, caption)

    # 4. Topic & Named Entities
    doc.add_heading("4. Topic & Named Entities", level=1)
    topic_imgs = [
        ("16_topic_fields_list.png", "Topic field list includes required fields."),
        ("17_topic_entities_student_details.png", "Named Entity: Student Details."),
        ("18_topic_entities_course_details.png", "Named Entity: Course Details."),
        ("19_topic_entities_prof_eval.png", "Named Entity: Professor Evaluation."),
        ("20_q_best_instructors_verified.png", "Verified Q answer: Best instructors."),
    ]
    for fn, caption in topic_imgs:
        insert_image_or_note(doc, fn, caption)

    # 5. Dashboard, Scenarios & Thread
    doc.add_heading("5. Dashboard, Scenarios & Thread", level=1)
    scenario_imgs = [
        ("23_publish_dashboard_q_enabled.png", "Student Enrollment Dashboard published with Q enabled."),
        ("25_scenario_create_select_visuals.png", "Scenario created with selected visuals."),
        ("26_scenario_starter.png", "Starter question: How do we improve professor evaluations while avoiding increased cost per course?"),
        ("34_scenario_thread_full.png", "Scenario thread showing iterative Q."),
    ]
    for fn, caption in scenario_imgs:
        insert_image_or_note(doc, fn, caption)

    # 6. Data Story
    doc.add_heading("6. Data Story", level=1)
    data_story_imgs = [
        ("36_data_story_cover_prepared_by.png", "Data Story cover with 'Prepared by'."),
        ("37_data_story_page1.png", "Data Story page showing thesis and supporting visuals."),
        ("38_data_story_page2.png", "Data Story recommendation and supporting visuals."),
    ]
    for fn, caption in data_story_imgs:
        insert_image_or_note(doc, fn, caption)

    # 7. File listing and artifacts
    doc.add_heading("7. File listings and additional artifacts", level=1)
    doc.add_paragraph("- dataset_fields.md\n- calculated_fields.md\n- Q_prompts.md\n- verified_answers.md\n- scenario_thread.md\n- data_story.md\n- manifest_sample.json\n- screenshots/ (all screenshot files present)")

    # 8. Methodology
    doc.add_heading("8. Methodology", level=1)
    doc.add_paragraph("Steps taken: sample dataset usage, attempted S3 import (error), created Q - Student Enrollment dataset from sample, added calculated field 'Student Type', created visuals using Amazon Q, configured Topic and Named Entities, created Scenario and Data Story.")

    # 9. Originality
    doc.add_heading("9. Originality Statement", level=1)
    doc.add_paragraph("I confirm this submission is my original work. Where official AWS/Udacity documentation was referenced, it is cited in the README.")

    # 10. Checklist for Rubric (simple placeholder)
    doc.add_heading("10. Checklist for Rubric", level=1)
    doc.add_paragraph("Map rubric sections to screenshots and files. (See README.md and this document's sections for details.)")

    # Save (overwrite submission_document.docx)
    outname = "submission_document.docx"
    doc.save(outname)
    print(f"Saved {outname}")

if __name__ == "__main__":
    main()
PY
